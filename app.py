import streamlit as st
import pandas as pd
import datetime
import ollama
from sqlalchemy import create_engine, Column, Integer, String, Text, DateTime, Enum
from sqlalchemy.ext.declarative import declarative_base
from sqlalchemy.orm import sessionmaker
from docx import Document
from io import BytesIO


# Setup the SQLite database connection and ORM model
DATABASE_URL = "sqlite:///incident_reports.db"
Base = declarative_base()

class IncidentReport(Base):
    __tablename__ = 'incident_reports'
   
    id = Column(Integer, primary_key=True)
    incident_type = Column(String(100))
    description = Column(Text)
    severity = Column(String(50))
    root_cause = Column(Text)
    corrective_action = Column(Text)
    assigned_reviewer = Column(String(100))
    date_reported = Column(DateTime, default=datetime.datetime.utcnow)
    status = Column(Enum('Open', 'Closed', name='incident_status'), default='Open')

# Database and session creation
engine = create_engine(DATABASE_URL, echo=True)
Base.metadata.create_all(engine)
Session = sessionmaker(bind=engine)
session = Session()

#  Ollama model prompting
def get_ollama_response(prompt, model="starling-lm:7b-alpha-q5_K_M"):
    try:
        response = ollama.chat(model=model, messages=[{"role": "user", "content": prompt}])
        return response['message']['content']
    except Exception as e:
        st.error(f"Error with Ollama API: {e}")
        return None

# Classify severity using Ollama model
def classify_severity(description):
    prompt = f"Classify the severity of this incident based on the following description: {description}. Classify as Critical, Major, or Minor."
    severity = get_ollama_response(prompt)
    if severity:
        return severity.strip()
    return "Unknown"

# Suggest root cause using Ollama model
def suggest_root_cause(description):
    prompt = f"Based on the following incident description, suggest the most likely root cause. Incident description: {description}"
    root_cause = get_ollama_response(prompt)
    if root_cause:
        return root_cause.strip()
    return "Unknown"

# Generate corrective action suggestion using Ollama model
def suggest_corrective_action(severity):
    prompt = f"Suggest corrective actions based on incident severity: {severity}. Provide a detailed corrective action plan."
    corrective_action = get_ollama_response(prompt)
    if corrective_action:
        return corrective_action.strip()
    return "Implement preventive measures."

# Generate the incident report
def generate_word_report(incident_type, description, severity, root_cause, corrective_action):
    doc = Document()
    doc.add_heading("Incident Report Summary", 0)
    
    doc.add_paragraph(f"Incident Type: {incident_type}")
    doc.add_paragraph(f"Description: {description}")
    doc.add_paragraph(f"Severity: {severity}")
    doc.add_paragraph(f"Root Cause: {root_cause}")
    doc.add_paragraph(f"Corrective Action: {corrective_action}")
    doc.add_paragraph(f"Date of Incident: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    
    # Save the document to a BytesIO buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# Streamlit UI for reporting an incident
st.title("Incident Reporting System")

st.header("Log a New Incident")

incident_type = st.selectbox("Select the Type of Incident", ["Accident", "Near Miss", "Safety Violation", "Equipment Failure", "Other"])

description = st.text_area("Describe the Incident", "")

assigned_reviewer = st.text_input("Assign a Reviewer (Safety Officer/Manager)", "")

if st.button("Generate Report"):
    if description:
        # Classify severity based on Ollama model
        severity = classify_severity(description)
       
        # Suggest root cause based on Ollama model
        root_cause = suggest_root_cause(description)
       
        # Suggest corrective action based on severity using Ollama model
        corrective_action = suggest_corrective_action(severity)
       
        # Generate the report
        report = generate_word_report(incident_type, description, severity, root_cause, corrective_action)
       
        # Show the incident report
        st.subheader("Incident Report Summary:")
        st.markdown(report)
       
        # Save incident data to the database
        new_report = IncidentReport(
            incident_type=incident_type,
            description=description,
            severity=severity,
            root_cause=root_cause,
            corrective_action=corrective_action,
            assigned_reviewer=assigned_reviewer,
            status='Open'
        )
       
        session.add(new_report)
        session.commit()
       
        # Optionally, allow the user to download the report
        st.download_button("Download Incident Report", report, file_name="incident_report.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    else:
        st.error("Please provide a description of the incident.")

# Review Incidents
st.header("Review Open Incidents")

open_incidents = session.query(IncidentReport).filter_by(status='Open').all()

incident_data = []
for incident in open_incidents:
    incident_data.append({
        "ID": incident.id,
        "Type": incident.incident_type,
        "Description": incident.description,
        "Severity": incident.severity,
        "Root Cause": incident.root_cause,
        "Assigned Reviewer": incident.assigned_reviewer,
        "Date Reported": incident.date_reported.strftime("%Y-%m-%d %H:%M:%S"),
        "Status": incident.status
    })

incident_df = pd.DataFrame(incident_data)
st.dataframe(incident_df)

# Option to mark incident as resolved
incident_id_to_close = st.number_input("Enter Incident ID to Close", min_value=1, max_value=len(open_incidents), step=1)

if st.button("Close Selected Incident"):
    if incident_id_to_close:
        incident_to_close = session.query(IncidentReport).filter_by(id=incident_id_to_close).first()
        if incident_to_close:
            incident_to_close.status = 'Closed'
            session.commit()
            st.success(f"Incident ID {incident_id_to_close} has been closed.")
        else:
            st.error("Incident ID not found.")
    else:
        st.error("Please select a valid incident ID.")
